"""Build and format the supplementary Word document for co-author review.

Usage:
    python paper1/scripts/build_supplementary_word_doc.py [supplementary_dir]
"""

import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

default_dir = Path(__file__).resolve().parents[1] / "manuscript"
supp_dir = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else default_dir
tex_path = supp_dir / "supplementary.tex"
raw_path = Path("/tmp/supplementary_raw.docx")
out_path = supp_dir / "supplementary.docx"

source = tex_path.read_text()
authors = re.findall(r"\\author\[([^]]+)\]\{([^}]+)\}", source)
affiliations = re.findall(r"\\affil\[([^]]+)\]\{([^}]+)\}", source)

subprocess.run(
    ["pandoc", str(tex_path), "-o", str(raw_path), "--from=latex"],
    check=True,
)

doc = Document(str(raw_path))

# Add affiliation indices to Pandoc's author paragraphs.
author_indices = {name.strip(): indices for indices, name in authors}
author_paragraphs = [p for p in doc.paragraphs if p.style.name == "Author"]
for paragraph in author_paragraphs:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = False
    indices = author_indices.get(paragraph.text.strip())
    if indices:
        paragraph.add_run(f" [{indices}]")

# Insert the affiliation block below the authors.
affiliation_nodes = set()
anchor = author_paragraphs[-1] if author_paragraphs else None
for indices, affiliation in affiliations:
    if anchor is None:
        break
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(f"{indices} {' '.join(affiliation.split())}")
    run.font.size = Pt(10)
    doc._body._body.remove(paragraph._p)
    anchor._p.addnext(paragraph._p)
    affiliation_nodes.add(paragraph._p)
    anchor = paragraph

title_paragraphs = [p for p in doc.paragraphs if p.style.name == "Title"]
for paragraph in title_paragraphs:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(12)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = True

if anchor is not None:
    anchor.add_run().add_break(WD_BREAK.PAGE)

# Restore supplementary table numbering in captions and cross-references.
caption_number = 0
for paragraph in doc.paragraphs:
    if paragraph.style.name == "Table Caption":
        caption_number += 1
        paragraph.text = f"Table S{caption_number}. {paragraph.text}"
        continue
    # Cross-reference text may sit inside a hyperlink and therefore may not
    # appear in ``paragraph.runs``. Update every Word text node directly.
    for text_node in paragraph._p.iter(qn("w:t")):
        for number in range(1, 8):
            text_node.text = (text_node.text or "").replace(
                f"Table\u00a0{number}", f"Table S{number}"
            )
            text_node.text = text_node.text.replace(
                f"Table {number}", f"Table S{number}"
            )
    for hyperlink in paragraph._p.iter(qn("w:hyperlink")):
        anchor_name = hyperlink.get(qn("w:anchor"), "")
        if not anchor_name.startswith("tab:s_"):
            continue
        for text_node in hyperlink.iter(qn("w:t")):
            if (text_node.text or "").isdigit():
                text_node.text = f"S{text_node.text}"

# Double-space prose. Keep titles, headings, affiliations, and captions single-spaced.
single_styles = {"Title", "Subtitle", "Author", "Table Caption"}
for paragraph in doc.paragraphs:
    style = paragraph.style.name or ""
    single_spaced = (
        style.startswith("Heading")
        or style in single_styles
        or paragraph._p in affiliation_nodes
    )
    paragraph.paragraph_format.line_spacing = 1.0 if single_spaced else 2.0
    if single_spaced:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if style == "Table Caption":
            paragraph.paragraph_format.keep_with_next = True
        continue
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.widow_control = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Keep table rows intact. Keep short tables on one page and repeat each header.
table_captions = [
    paragraph for paragraph in doc.paragraphs if paragraph.style.name == "Table Caption"
]


def border_element(tag):
    border = OxmlElement(f"w:{tag}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "8")
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), "000000")
    return border


def insert_before_first(parent, element, later_tags):
    later_qnames = {qn(f"w:{tag}") for tag in later_tags}
    for index, child in enumerate(parent):
        if child.tag in later_qnames:
            parent.insert(index, element)
            return
    parent.append(element)


for table_index, table in enumerate(doc.tables):
    tbl_pr = table._tbl.tblPr
    for old in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(old)
    tbl_borders = OxmlElement("w:tblBorders")
    for tag in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tbl_borders.append(border_element(tag))
    insert_before_first(
        tbl_pr,
        tbl_borders,
        ("shd", "tblLayout", "tblCellMar", "tblLook", "tblCaption", "tblDescription", "tblPrChange"),
    )

    keep_table_together = len(table.rows) <= 12
    if 8 <= len(table.rows) <= 12:
        table_captions[table_index].paragraph_format.page_break_before = True
    final_row_index = len(table.rows) - 1
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        if row_index == 0 and tr_pr.find(qn("w:tblHeader")) is None:
            tr_pr.append(OxmlElement("w:tblHeader"))
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.0
                if keep_table_together and row_index < final_row_index:
                    paragraph.paragraph_format.keep_with_next = True

# Add continuous line numbering to every section.
for section in doc.sections:
    sect_pr = section._sectPr
    for old in sect_pr.findall(qn("w:lnNumType")):
        sect_pr.remove(old)
    line_numbers = OxmlElement("w:lnNumType")
    line_numbers.set(qn("w:countBy"), "1")
    line_numbers.set(qn("w:start"), "1")
    line_numbers.set(qn("w:restart"), "continuous")
    line_numbers.set(qn("w:distance"), "360")
    sect_pr.append(line_numbers)

doc.save(str(out_path))
print(
    f"Wrote {out_path} "
    f"({len(doc.tables)} tables; double-spaced prose; "
    "single-spaced tables and captions; continuous line numbers)"
)
