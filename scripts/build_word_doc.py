"""
Build a PeerJ-formatted Word version of manuscript.tex. Uses pandoc with
citeproc to render headings,
bold/italic, citations as author-year, and a full reference list. All
figures, tables, and the algorithm block are replaced with text
placeholders that show the caption and the LaTeX label.

Usage:
    python paper1/scripts/build_word_doc.py

Outputs:
    paper1/latex/peerj/manuscript.docx     (updated in place)

Requires:
    pandoc (with citeproc)
"""

import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx_omml_repair import repair_pandoc_omml

_default_dir = Path(__file__).resolve().parents[1] / "manuscript"
_dir = sys.argv[1] if len(sys.argv) > 1 else _default_dir
src_path = Path(_dir) / "manuscript.tex"
out_path = Path("/tmp/manuscript_for_word.tex")

src = src_path.read_text()

# Preserve the title-page metadata. Pandoc's LaTeX reader does not render the
# optional affiliation arguments used by authblk, so the Word post-processing
# step inserts a compact title and affiliation block explicitly.
title_match = re.search(r"\\title\{(.+?)\}", src, flags=re.DOTALL)
word_title = re.sub(r"\s+", " ", title_match.group(1)).strip() if title_match else ""
word_authors = [
    (indices, name.strip())
    for indices, name in re.findall(r"\\author\[([^]]+)\]\{([^}]+)\}", src)
]
word_affils = [
    (indices, re.sub(r"\s+", " ", text).strip())
    for indices, text in re.findall(r"\\affil\[([^]]+)\]\{([^}]+)\}", src)
    if indices != "*"
]
corresponding_match = re.search(
    r"\\corrauthor\[([^]]+)\]\{([^}]+)\}\{([^}]+)\}",
    src,
)
word_corresponding = (
    tuple(part.strip() for part in corresponding_match.groups())
    if corresponding_match
    else None
)

# Remove the lineno option from the documentclass
src = src.replace("[fleqn,10pt,lineno]", "[fleqn,10pt]")

# Preserve simple signed values such as $+0.4$ in Word. Pandoc can interpret
# the leading plus as an empty unary expression and omit the number.
src = re.sub(
    r"\$([+-])\s*([0-9]+(?:\.[0-9]+)?)\$",
    lambda m: m.group(1) + m.group(2),
    src,
)

# Drop algorithm2e and placeins packages (pandoc chokes on algorithm2e)
src = re.sub(r"\\usepackage\[ruled,linesnumbered,vlined\]\{algorithm2e\}\s*\n", "", src)
src = re.sub(r"\\usepackage\[section\]\{placeins\}\s*\n", "", src)
src = re.sub(r"\\usepackage\{needspace\}\s*\n", "", src)

# Remove \needspace, \FloatBarrier, \clearpage commands
src = re.sub(r"\\needspace\{[^}]*\}\s*\n?", "", src)
src = re.sub(r"\\FloatBarrier\s*\n?", "", src)
src = re.sub(r"\\clearpage\s*\n?", "", src)

# Extract abstract content and move it inside the document body
abstract_match = re.search(
    r"\\begin\{abstract\}\s*(.+?)\s*\\end\{abstract\}",
    src,
    flags=re.DOTALL,
)
abstract_body = ""
if abstract_match:
    abstract_body = abstract_match.group(1)
    # Remove from its original position (which is in the preamble)
    src = src[: abstract_match.start()] + src[abstract_match.end() :]
    # Insert after \maketitle so it appears at the top of the body
    src = src.replace(
        r"\maketitle",
        r"\maketitle" + "\n\n\\section*{Abstract}\n" + abstract_body + "\n",
        1,
    )


# Helper: find balanced block from \begin{env} to matching \end{env}
def replace_env(text, env, repl_fn):
    """Replace each \\begin{env}...\\end{env} block using repl_fn(body)."""
    pattern = re.compile(
        r"\\begin\{" + env + r"\*?\}(.*?)\\end\{" + env + r"\*?\}",
        re.DOTALL,
    )
    return pattern.sub(lambda m: repl_fn(m.group(1)), text)


def extract_caption(body):
    """Extract caption text from a figure/table body. Returns plain text."""
    m = re.search(r"\\caption\{(.+?)\}\s*\n", body, re.DOTALL)
    if not m:
        # Try non-newline-terminated
        m = re.search(r"\\caption\{(.+?)\}", body, re.DOTALL)
    if not m:
        return "(no caption found)"
    # Strip outer braces and clean
    caption = m.group(1)
    # Crude: remove common LaTeX commands
    caption = re.sub(r"\\textbf\{([^}]+)\}", r"**\1**", caption)
    caption = re.sub(r"\\textit\{([^}]+)\}", r"*\1*", caption)
    caption = re.sub(r"\\emph\{([^}]+)\}", r"*\1*", caption)
    caption = re.sub(r"\\cite\w*\{[^}]+\}", "", caption)
    caption = re.sub(r"\\ref\{[^}]+\}", "?", caption)
    caption = re.sub(r"\\label\{[^}]+\}", "", caption)
    caption = caption.replace("~", " ")
    caption = re.sub(r"\s+", " ", caption).strip()
    return caption


def extract_label(body):
    m = re.search(r"\\label\{([^}]+)\}", body)
    return m.group(1) if m else ""


# Resolve numbered display references before replacing the LaTeX environments.
# Pandoc otherwise renders \ref{tab:...} literally as [tab:...].
display_refs = {}
for env, prefix in [("figure", "fig:"), ("table", "tab:"), ("algorithm", "alg:")]:
    pattern = re.compile(
        r"\\begin\{" + env + r"\*?\}(.*?)\\end\{" + env + r"\*?\}",
        re.DOTALL,
    )
    number = 0
    for match in pattern.finditer(src):
        number += 1
        label = extract_label(match.group(1))
        if label:
            display_refs[label] = str(number)

src = re.sub(
    r"\\ref\{((?:fig|tab|alg):[^}]+)\}",
    lambda match: display_refs.get(match.group(1), match.group(0)),
    src,
)


fig_counter = [0]
tab_counter = [0]


def figure_placeholder(body):
    fig_counter[0] += 1
    caption = extract_caption(body)
    label = extract_label(body)
    n = fig_counter[0]
    return (
        "\n\n\\begin{quote}\n"
        f"\\textbf{{[FIGURE {n}: placeholder. Label: \\texttt{{{label}}}]}}\n\n"
        f"\\textit{{Caption: {caption}}}\n"
        "\\end{quote}\n\n"
    )


def table_placeholder(body):
    tab_counter[0] += 1
    caption = extract_caption(body)
    label = extract_label(body)
    n = tab_counter[0]
    return (
        "\n\n\\begin{quote}\n"
        f"\\textbf{{[TABLE {n}: placeholder. Label: \\texttt{{{label}}}]}}\n\n"
        f"\\textit{{Caption: {caption}}}\n"
        "\\end{quote}\n\n"
    )


def algorithm_placeholder(body):
    return (
        "\n\n\\begin{quote}\n"
        "\\textbf{[ALGORITHM 1: placeholder. Bilateral ZUPT stride-length "
        "estimation pseudocode]}\n\n"
        "\\textit{See manuscript PDF for the full pseudocode listing.}\n"
        "\\end{quote}\n\n"
    )


src = replace_env(src, "figure", figure_placeholder)
src = replace_env(src, "table", table_placeholder)
src = replace_env(src, "algorithm", algorithm_placeholder)

# Replace remaining \multirow{N}{*}{...} with just the inner content
src = re.sub(r"\\multirow\{\d+\}\{\*\}\{([^}]*)\}", r"\1", src)

# Pandoc dislikes the \dps macro we defined; remove the definition and expand usages
src = re.sub(r"\\newcommand\{\\dps\}\{[^}]*\}\s*\n?", "", src)
src = re.sub(r"\\dps(?![a-zA-Z])", "°/s", src)

out_path.write_text(src)
print(f"Wrote {out_path} ({len(src)} chars)")
print(f"Replaced {fig_counter[0]} figures, {tab_counter[0]} tables")

# ---------------------------------------------------------------------------
# Render to a PeerJ-formatted .docx via pandoc.
# ---------------------------------------------------------------------------
peerj_dir = src_path.parent
bib_path = peerj_dir / "references.bib"
docx_out = peerj_dir / "manuscript.docx"
docx_tmp = Path("/tmp/manuscript_raw.docx")

subprocess.run(
    [
        "pandoc",
        str(out_path),
        "--citeproc",
        "--bibliography",
        str(bib_path),
        "--metadata",
        "reference-section-title=References",
        "-o",
        str(docx_tmp),
    ],
    check=True,
)

# Format the review copy: justify and double-space the body, keep display
# material single-spaced, and add continuous line numbers.
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    doc = Document(str(docx_tmp))

    # Pandoc may drop all authblk title-page metadata. Restore it when needed,
    # then insert affiliations immediately after the author line.
    author_paragraphs = [p for p in doc.paragraphs if p.style.name == "Author"]
    if not author_paragraphs and word_title and word_authors:
        first = doc.paragraphs[0]
        title_p = doc.add_paragraph(style="Title")
        title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_p.add_run(word_title)
        doc._body._body.remove(title_p._p)
        first._p.addprevious(title_p._p)

        author_p = doc.add_paragraph(style="Author")
        author_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        author_p.add_run(
            ", ".join(f"{name} [{indices}]" for indices, name in word_authors)
        )
        doc._body._body.remove(author_p._p)
        title_p._p.addnext(author_p._p)
        author_paragraphs = [author_p]

    for p in doc.paragraphs:
        if p.style.name in ("Title", "Author"):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    author_indices = {name: indices for indices, name in word_authors}
    for p in author_paragraphs:
        indices = author_indices.get(p.text.strip())
        if indices:
            p.add_run(f" [{indices}]")
    anchor = author_paragraphs[-1] if author_paragraphs else None
    affiliation_nodes = set()
    for indices, text in word_affils:
        if not text or anchor is None:
            continue
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"{indices} {text}")
        run.font.size = Pt(10)
        doc._body._body.remove(p._p)
        anchor._p.addnext(p._p)
        affiliation_nodes.add(p._p)
        anchor = p

    if word_corresponding and anchor is not None:
        corr_indices, corr_name, corr_email = word_corresponding
        corr_address = next(
            (text for indices, text in word_affils if indices == corr_indices),
            "",
        )
        correspondence_lines = [
            "Corresponding Author",
            f"{corr_name} [{corr_indices}]",
            corr_address,
            f"Email address: {corr_email}",
        ]
        for index, line in enumerate(correspondence_lines):
            if not line:
                continue
            p = doc.add_paragraph(style="Normal")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            run.font.size = Pt(10)
            if index == len(correspondence_lines) - 1:
                run.add_break(WD_BREAK.PAGE)
            doc._body._body.remove(p._p)
            anchor._p.addnext(p._p)
            affiliation_nodes.add(p._p)
            anchor = p

    for style in doc.styles:
        if style.type == 1:
            style.font.name = "Times New Roman"
            style._element.get_or_add_rPr().get_or_add_rFonts().set(
                qn("w:eastAsia"), "Times New Roman"
            )
    doc.styles["Normal"].font.size = Pt(12)

    skip = ("Title", "Subtitle", "Author")
    n = 0
    for p in doc.paragraphs:
        name = p.style.name or ""
        text = p.text.strip()
        single_spaced = (
            name.startswith("Heading")
            or name in skip
            or p._p in affiliation_nodes
            or name == "Block Text"
            or text.startswith("Caption:")
            or bool(p._p.xpath(".//m:oMathPara"))
        )
        p.paragraph_format.line_spacing = 1.0 if single_spaced else 2.0
        if single_spaced:
            if not p._p.xpath(".//w:drawing") and not p._p.xpath(".//m:oMathPara"):
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if name == "Block Text":
                p.paragraph_format.keep_together = True
                if text.startswith(("[FIGURE", "[TABLE", "[ALGORITHM")):
                    p.paragraph_format.keep_with_next = True
            continue
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.widow_control = True
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        n += 1

    # Remove any text highlighting or paragraph/run shading in the abstract.
    # Pandoc does not normally add it, but this makes the Word output match the
    # white-background abstract in the PDF if a reference style is introduced.
    in_abstract = False
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            in_abstract = p.text.strip() == "Abstract"
            continue
        if not in_abstract:
            continue
        p_pr = p._p.get_or_add_pPr()
        for shd in p_pr.findall(qn("w:shd")):
            p_pr.remove(shd)
        for run in p.runs:
            run.font.highlight_color = None
            r_pr = run._r.get_or_add_rPr()
            for shd in r_pr.findall(qn("w:shd")):
                r_pr.remove(shd)

    # Add continuous line numbering to every document section.
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)

        sect_pr = section._sectPr
        page_margins = sect_pr.find(qn("w:pgMar"))
        page_margins.set(qn("w:header"), "720")
        page_margins.set(qn("w:footer"), "720")
        page_margins.set(qn("w:gutter"), "0")
        for old in sect_pr.findall(qn("w:lnNumType")):
            sect_pr.remove(old)
        line_numbers = OxmlElement("w:lnNumType")
        line_numbers.set(qn("w:countBy"), "1")
        line_numbers.set(qn("w:start"), "1")
        line_numbers.set(qn("w:restart"), "continuous")
        line_numbers.set(qn("w:distance"), "360")
        sect_pr.append(line_numbers)

        footer_paragraph = section.footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_ppr = footer_paragraph._p.get_or_add_pPr()
        suppress_line_number = OxmlElement("w:suppressLineNumbers")
        footer_justification = footer_ppr.find(qn("w:jc"))
        footer_ppr.insert(
            footer_ppr.index(footer_justification),
            suppress_line_number,
        )
        page_run = footer_paragraph.add_run()
        field_begin = OxmlElement("w:fldChar")
        field_begin.set(qn("w:fldCharType"), "begin")
        field_code = OxmlElement("w:instrText")
        field_code.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        field_code.text = " PAGE "
        field_end = OxmlElement("w:fldChar")
        field_end.set(qn("w:fldCharType"), "end")
        page_run._r.extend((field_begin, field_code, field_end))

    doc.save(str(docx_out))
    print(
        f"Wrote {docx_out} "
        f"(PeerJ letter layout; justified {n} body paragraphs; "
        "continuous line and page numbers)"
    )
except ImportError:
    shutil.copy(docx_tmp, docx_out)
    print(f"Wrote {docx_out} (python-docx unavailable; not justified)")

omml_repairs = repair_pandoc_omml(docx_out)
if omml_repairs:
    print(f"Repaired {omml_repairs} invalid Pandoc OMML normal-math runs")
